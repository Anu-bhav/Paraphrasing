#!/usr/bin/env python3
"""
Author: Anubhav
Date: 29.01.2020

The program will take a processed textfile as input, the file should have every sentence in a new line

Login info to get cookies:
Url: https://quillbot.com/
Email: prinusam.roy.9e@819760.com
Password: password
"""
import os.path
from json import loads
from urllib.parse import quote
from docx import Document
import requests

API_URL = "https://quillbot.com/api/singleParaphrase"
PARAMS = "?userID=N/A&text={}&strength={}&autoflip={}&wikify={}&fthresh={}"
DOC = Document()


def setup_session():
    """Update headers for the session.

    Returns
    -------
    obj
        Requests Session
    """
    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": "Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:68.0) Gecko/20100101 Firefox/68.0",
            "Content-Type": "application/text"
        }
    )
    return session


def get_parameterized_url(text):
    """Gets parametrized url

    Parameters
    ----------
    text : str

    Returns
    -------
    url: str
        string containing url and quoted text
    """
    url_encoded_text = quote(text)
    autoflip = "true"
    fthresh = "4"
    strength = "5"
    wikify = "9"
    url = API_URL + PARAMS.format(url_encoded_text, strength, autoflip, wikify, fthresh)
    return url


def openfile(filepath):
    """open the txt file to be process and saves it to a docx file

    Arguments:
        filepath {string} -- the path, that the user inputs
    """

    filepath = filepath.rstrip()
    filesave = os.path.split(filepath)[-1]
    filesave = "Drafts/" + ([filename.strip() for filename in filesave.split('.')][-2] + " - draft")

    cur_path = os.path.dirname(__file__)
    textfile = os.path.relpath(filepath, cur_path)

    with open(textfile, "r") as wholetext:
        for line in wholetext:
            line = line.rstrip()

            session = setup_session()
            if len(line) > 700:
                print("line should be less than 700 characters, %s" % line)
            else:
                url = get_parameterized_url(line)
                paraphrasor(url, session)
                DOC.save('%s.docx' % filesave)
    wholetext.close()


def paraphrasor(url, session):
    """Gets paraphrased text

    Parameters
    ----------
    url : str
        Complete url containing text
    session : class `requests.sessions.Session`
            Requests session.
            Provides cookie persistence, connection-pooling, and configuration.
    """
    # Cookies are configurable
    cookies = {
        "__cfduid": "dc0840a2a6813fdf6a4cdc61e6f9fccde1581242299;",
        "_ga": "GA1.2.2004764009.1581242301;",
        "_gid": "GA1.2.19653033.1581242301;",
        "connect.sid": "s%3Ahe1zGWUi5utvCE_O-PEh3Rb3moDifGYy.1kQ1L8iIO5BaBTtyi9re8jFZIG2YBmMO14fbTwxGYto;",
        "userIDToken": "eyJhbGciOiJSUzI1NiIsImtpZCI6IjYzZTllYThmNzNkZWExMTRkZWI5YTY0OTcxZDJhMjkzN2QwYzY3YWEiLCJ0eXAiOiJKV1QifQ.eyJpc3MiOiJodHRwczovL3NlY3VyZXRva2VuLmdvb2dsZS5jb20vcGFyYXBocmFzZXItNDcyYzEiLCJhdWQiOiJwYXJhcGhyYXNlci00NzJjMSIsImF1dGhfdGltZSI6MTU4MTI0MjMxOCwidXNlcl9pZCI6IlJxRUhUZUxEQlBjWlRzSTFyWW81TWZjdTNkMDIiLCJzdWIiOiJScUVIVGVMREJQY1pUc0kxcllvNU1mY3UzZDAyIiwiaWF0IjoxNTgxMjQyMzIwLCJleHAiOjE1ODEyNDU5MjAsImVtYWlsIjoicHJpbnVzYW0ucm95LjllQDgxOTc2MC5jb20iLCJlbWFpbF92ZXJpZmllZCI6dHJ1ZSwiZmlyZWJhc2UiOnsiaWRlbnRpdGllcyI6eyJlbWFpbCI6WyJwcmludXNhbS5yb3kuOWVAODE5NzYwLmNvbSJdfSwic2lnbl9pbl9wcm92aWRlciI6InBhc3N3b3JkIn19.dBIarlzbCSdZ4GdOoUy8sB_8jzIGdccj3oz0L6GFQkI5yCUWhgJs5XXMESV2QM4QamPm5MPstqLV2Tc1CGRijADBixMS5ZwqVk7twLTmriqZVYwrFEEhtsZYi3VwBLwcetRt-9jrBnzrFK0chus8_g7tf--HPTqn4VN4dXzu70GGMz6nbUVf7_dOmlJReUUWw5Ff2qorMMOYyh6_UhdrCQBcKz3bX_5WJBUPr5vR4cxDyIdAft9Idizi8-5tMJ2UwlTJ-D0DQviN2lUWJ_kJaKoL9dAuQgkryeGBAGLzxIdZlhCPZpaaro8t32QpYgrUWTu0j4PjiRZbM8hu0NH6Ag;",
        "authenticated": "true;",
        "premium": "false;",
        "quid": "RqEHTeLDBPcZTsI1rYo5Mfcu3d02;",
        "REGION": "uk;",
        "amplitude_id_6e403e775d1f5921fdc52daf8f866c66quillbot.com": "eyJkZXZpY2VJZCI6IjgwM2YwMGQ4LTc1ODgtNDBmMy1iMzczLWVkM2NkZDVjYTU5OVIiLCJ1c2VySWQiOiJwcmludXNhbS5yb3kuOWVAODE5NzYwLmNvbSIsIm9wdE91dCI6ZmFsc2UsInNlc3Npb25JZCI6MTU4MTI2MTc0MDkxOCwibGFzdEV2ZW50VGltZSI6MTU4MTI2MzcxNTkzNCwiZXZlbnRJZCI6MzAxLCJpZGVudGlmeUlkIjowLCJzZXF1ZW5jZU51bWJlciI6MzAxfQ==;",
        "_gat": "1;",
        "sessID": "ff9e579bb9ef1a98"
    }

    req = session.get(url, cookies=cookies)
    if req.status_code == 200:
        json_text = loads(req.text)

        end = "\n\n"

        json_text = json_text[0] if len(json_text) == 1 else json_text
        print(f"\nData Sent: {json_text['sent']}", end)
        DOC.add_paragraph(f"\nData Sent: {json_text['sent']}")

        paras = [key for key in json_text if key.startswith("paras")]
        texts = list(
            {text.get("alt") for para in paras for text in json_text[para]}
        )

        print("Alternative Texts:")
        DOC.add_paragraph("Alternative Texts:")

        print("_" * 90, "\n")
        DOC.add_paragraph("_" * 105)
        # DOC.add_paragraph()

        for text in texts:
            print(text, end)
            DOC.add_paragraph(text)

        print("#" * 90)
        DOC.add_paragraph("#" * 63)
        DOC.add_page_break()


def main():
    """main function the program starts
    """
    print("Quillbot Paraphrasing Tool")
    filepath = input("Enter the path of the formatted text file: ")
    openfile(filepath)


if __name__ == "__main__":
    main()
